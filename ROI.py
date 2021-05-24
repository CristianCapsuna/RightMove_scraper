######## COMMENTS #######
# 1) This is an old program I wrote. It needs updating as rightmove have updated their website but there were also a few drawbacks with it that could use improving.

###### IMPROVEMENTS #####
# 1) Need to find another way to mine the data. Using this kind of predictible behaviour will get you banned by the website anti-sraping software.
# Rightmove have an API that can be used but also I have found out Selenium is not the right tool for the job. I am sure there is a more efficient scraper
# that can scrape only the required by and not the entire website.
# 2) On the data side, the next step would be to create a SQL databse to store the data so that trends accross time can be determined.
#

from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from datetime import datetime,timedelta
from time import sleep
from threading import Lock, Thread
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import TimeoutException
from selenium.common.exceptions import ElementClickInterceptedException
from selenium.common.exceptions import WebDriverException
import os
from pathlib import Path
from selenium.webdriver.common.desired_capabilities import DesiredCapabilities
import docx
from docx.enum.dml import MSO_THEME_COLOR_INDEX

##########################   ASSUMPTIONS ##########################
############Leicestr#Doncstr#Manchstr#Liverpl#Biham#Aberdn#Sheffield#
reg_codes = ["5E789","5E430","5E904","5E813","5E162","5E4","5E1195"]
maxprice = 100000
minprice = 60000
target_ROI = 18    #percentage
BRRR = False #False or True
ignore_string = ["non-standard construction".lower(),"Lease hold".lower(),"Leasehold".lower()]
separate_file_string = [["auction".lower(),"bidding".lower(),"reserve price".lower()],["cash buy".lower()]]
work_path = "C:/Users/Cristian/Desktop/ROI_work_area/"
###################################################################

#####Parallel thread which takes care of calculating the estimated progress of the program.
def progress_bar():
    global stop_marker
    global current_page_number
    global total_number_of_pages
    global lock
    global number_of_offers_for_page
    global current_offer_in_page_increment
    global time_per_offer
    global offer_counter
    global number_of_regions
    base_percentage = 0
    detail_percentage = 0
    current_progress = 0
    time_per_offer = 0
    time_per_offer_measurement_counter = 1
    sum_of_time_measurements = timedelta()
    print_marker = False
    ######This is to find out how much each page is of the total number of pages in percentage
    lock.acquire()
    while True:
        try:
            percentage_per_page = 100 / total_number_of_pages
        except ZeroDivisionError:
            continue
        else: break
    lock.release()
    ######This is to estimate the total number of offers that are going to be looked at by the program so it can estimate the time it will take
    estimated_number_of_offers = (total_number_of_pages - number_of_regions) * 25 + number_of_regions * 12
    print("Estimated number of offers: " + str(estimated_number_of_offers))
    while not stop_marker:
        lock.acquire()
        #####The calculates the percentage of the total job, per offer for a particular page with the below caveat
        #####Every page, irrespective of how many offers it contains is considered by the program to be the same percentage of the total job. Not ideal but will do for now
        percentage_per_offer = percentage_per_page / number_of_offers_for_page
        #####Checks if a new page has been started. If a new page has been started the base percentage is adjusted and the detail one set to 0 asa you are now on the first offer per page
        if current_page_number * percentage_per_page > base_percentage:
            base_percentage = current_page_number * percentage_per_page
            detail_percentage = 0
            lock.release()
            if int(base_percentage) > int(current_progress):
                print_marker = True
            #####Variable that tracks the current progress
            current_progress = base_percentage
            if print_marker:
                print(str(int(current_progress)) + "% of search finished" + ", base_percentage = " + str(base_percentage))
                print_marker = False
            sleep(3)
        #####If a new page has not been started then it it will just calculate the progress from old base percentage
        elif percentage_per_offer * current_offer_in_page_increment > detail_percentage:
            detail_percentage = percentage_per_offer * current_offer_in_page_increment
            lock.release()
            if int(base_percentage + detail_percentage) > int(current_progress):
                print_marker = True
            current_progress = base_percentage + detail_percentage
            if print_marker:
                print(str(int(current_progress)) + "% of search finished" + ", base_percentage = " + str(base_percentage) + ", detail_percentage = " + str(detail_percentage))
                print_marker = False
                #####This section calculates the estimated remaining time and estimated completion time based on it.
                try:
                    sum_of_time_measurements = sum_of_time_measurements + time_per_offer
                    completion_time = (estimated_number_of_offers - offer_counter) * (sum_of_time_measurements / time_per_offer_measurement_counter)
                    time_per_offer_measurement_counter += 1
                    if completion_time != 0:
                        print("Estimated time left: " + str(completion_time) + ", Estimated completion time: " + str(datetime.now() + completion_time))
                except NameError:
                    pass
            sleep(3)
        else:
            lock.release()
            sleep(3)
         

class ratio_calculator():
    def __init__(self,supply_link,supply_and_demand_link):
        self.supply = self.worker(supply_link)
        self.supply_and_demand = self.worker(supply_and_demand_link)
        try:
            self.ratio = self.supply/(self.supply_and_demand-self.supply)
        except ZeroDivisionError:
            self.ratio = "No demand for"
    def worker(self,link):
        while True:
            try:
                driver.get(link)
                WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.ID, 'searchHeader')))
                result_string = driver.find_element_by_id('searchHeader').text.split()[0]
                if "," in result_string:
                    result_string = result_string.replace(",", "")
                return float(result_string)
            except TimeoutException:
                try:
                    if driver.find_element_by_xpath('//p[@class = "enhancedZeroResults-title"]').text :
                        return float(0)
                    break
                except NoSuchElementException:
                    driver.refresh()
            else:
                break

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

####This tells chrome to not wait until the page is fully loaded
Page_load_strategy = DesiredCapabilities.CHROME
Page_load_strategy["pageLoadStrategy"] = "none"
#####These are for measuring the search duration on number of offered processed by the program
offer_counter = 0
begin_time = datetime.now()
#####Defining and initializing values needed for the progress analysis of the program by the parallel thread
lock = Lock()
stop_marker = False
current_page_number = 0
total_number_of_pages = 0
number_of_offers_for_page = 1
current_offer_in_page_increment = 0

driver = webdriver.Chrome("C:/Users/Cristian/AppData/Local/Programs/Python/ROI_calculator/ChromeDriver/chromedriver.exe", desired_capabilities=Page_load_strategy)
driver.maximize_window()

#####Checks to see if the files it will write to already exist and if they do overwrites them
work_area_path = Path(work_path)
files = os.listdir(work_area_path)
for file in files:
    os.remove(Path(work_area_path, file))

trailing_list = []
for i in range(1, len(separate_file_string)+2):
    trailing_list.append(0)

number_of_regions = len(reg_codes)

#####The bit below will check how many offer there are to go through
for code in reg_codes:
    driver.get("https://www.rightmove.co.uk/property-for-sale/find.html?locationIdentifier=REGION%" + code + "&minBedrooms=2&maxPrice=" + str(maxprice) + "&minPrice="+ str(minprice) +"&propertyTypes=detached" + \
        "%2Csemi-detached%2Cterraced&primaryDisplayPropertyType=houses&includeSSTC=false&mustHave=&dontShow=sharedOwnership%2Cretirement&furnishTypes=&keywords=")
    while True:
        try:
            WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//h1[@class = "searchTitle-heading"]')))
            no_of_pages = driver.find_elements_by_xpath('//span[@class = "pagination-pageInfo"]')[-1].text
            if no_of_pages == "":
                break
            lock.acquire()
            total_number_of_pages += int(no_of_pages)
            lock.release()
        except TimeoutException:
            driver.refresh()
        except IndexError:
            driver.refresh()
        else: break

#####Start the thread meant to monitor the progress
thread = Thread(target=progress_bar)
thread.start()

######Master loop to go through every one of the region codes provided
for code in reg_codes:
    for i in range(0,len(trailing_list)):
        trailing_list[i] = 0
    #####Gets the supply/demand for sales of houses similars to the ones we are seeking
    sales_ratio = ratio_calculator("https://www.rightmove.co.uk/property-for-sale/find.html?locationIdentifier=REGION%" + code + "&maxBedrooms=3&minBedrooms=2&maxPrice=" + str(maxprice) + "&propertyTypes=detached" + \
        "%2Csemi-detached%2Cterraced&primaryDisplayPropertyType=houses&includeSSTC=false&mustHave=&dontShow=sharedOwnership%2Cretirement&furnishTypes=&keywords=", \
        "https://www.rightmove.co.uk/property-for-sale/find.html?locationIdentifier=REGION%" + code + "&maxBedrooms=3&minBedrooms=2&maxPrice=" + str(maxprice) + "&propertyTypes=detached" + \
        "%2Csemi-detached%2Cterraced&primaryDisplayPropertyType=houses&includeSSTC=true&mustHave=&dontShow=sharedOwnership%2Cretirement&furnishTypes=&keywords=").ratio
    #####Gets the supply/demand for renting of houses similars to the ones we are seeking
    rent_ratio = ratio_calculator("https://www.rightmove.co.uk/property-to-rent/find.html?locationIdentifier=REGION%" + code + "&maxBedrooms=3&minBedrooms=2&propertyTypes=detached%2Csemi-detached%2Cterraced" + \
        "&primaryDisplayPropertyType=houses&includeLetAgreed=false&mustHave=&dontShow=retirement&furnishTypes=&keywords=", \
        "https://www.rightmove.co.uk/property-to-rent/find.html?locationIdentifier=REGION%" + code + "&maxBedrooms=3&minBedrooms=2&propertyTypes=detached%2Csemi-detached%2Cterraced" + \
        "&primaryDisplayPropertyType=houses&includeLetAgreed=true&mustHave=&dontShow=retirement&furnishTypes=&keywords=").ratio

    #####Generated the seach to be used for that region code.
    driver.get("https://www.rightmove.co.uk/property-for-sale/find.html?locationIdentifier=REGION%" + code + "&minBedrooms=2&maxPrice=" + str(maxprice) + "&propertyTypes=detached" + \
               "%2Csemi-detached%2Cterraced&primaryDisplayPropertyType=houses&includeSSTC=false&mustHave=&dontShow=sharedOwnership%2Cretirement&furnishTypes=&keywords=")

    while True:
        try:
            area_of_search = WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//h1[@class = "searchTitle-heading"]'))).text.split(", up to ")[0]
        except TimeoutException:
            driver.refresh()
        else:
            break

    #####Finds the number of pages that a region search has
    no_of_pages = driver.find_elements_by_xpath('//span[@class = "pagination-pageInfo"]')[-1].text
    #####If there are no pages because there are no properties it will move on to the next region
    if no_of_pages == "":
        continue
    #####Second level loop which goes through all pages of results for a region
    for i in range(1,int(no_of_pages)+1):
        lock.acquire()
        current_offer_in_page_increment = 0
        lock.release()
        #####For all the page except the first apge it will go back to the search to go to the needed page
        while True:
            try:
                if i>1:
                    driver.get("https://www.rightmove.co.uk/property-for-sale/find.html?locationIdentifier=REGION%" + code + "&minBedrooms=2&maxPrice=" + str(maxprice) + "&minPrice="+ str(minprice) +"&propertyTypes=detached" + \
        "%2Csemi-detached%2Cterraced&primaryDisplayPropertyType=houses&includeSSTC=false&mustHave=&dontShow=sharedOwnership%2Cretirement&furnishTypes=&keywords=")
                WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//select[@class = "select pagination-dropdown"]//option')))
                page_options = driver.find_elements_by_xpath('//select[@class = "select pagination-dropdown"]//option')
            except TimeoutException:
                driver.refresh()
            else:
                break
        #####Clicks the required page number in the sequence
        if i != 1:
            page_options[i-1].click()
        #####Grabs all the offer from the page
        while True:
            try:
                WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//a[@class = "propertyCard-priceLink propertyCard-salePrice"]')))
                ads_list = driver.find_elements_by_xpath('//a[@class = "propertyCard-priceLink propertyCard-salePrice"]')
            except TimeoutException:
                driver.refresh()
            else:
                break
        #####Collates all the URL links of the offers on the page
        URL_list_of_property_for_sale = []
        for element in ads_list:
            URL_list_of_property_for_sale = URL_list_of_property_for_sale + [element.get_attribute('href')]
        #####Grabs the number of offers on the page for the proress tracking function
        lock.acquire()
        number_of_offers_for_page = len(URL_list_of_property_for_sale)
        lock.release()

        current_window_handle = driver.current_window_handle
        #####Lowest level loop going through all the offers on the page
        for offer_URL in URL_list_of_property_for_sale:
            output_file_picker = ""
            #####Calculation of time spent per offer needed for estimating remaining time by the progress function.
            lock.acquire()
            try:
                time_per_offer = datetime.now() - offer_start_time
            except NameError:
                pass
            lock.release()
            offer_start_time = datetime.now()
            while True:
                try:
                    #####Counter to count offers processed
                    offer_counter = offer_counter + 1
                    #####Break marker used to break the loop if a property has the selling price set to "POA"
                    break_marker = False
                    #####Accessing the URL of the offer
                    driver.get(offer_URL)
                    while True:
                        try:
                            #####Section probest the property description for keywords for filtering. Work in progress
                            try:
                                total_string = ''
                                # WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//div[@class = "left overflow-hidden agent-content"]//div[@class = "sect "]')))
                                WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//div[@class = "left overflow-hidden agent-content"]//div[@class = "sect "]//div[@class = "sect "]|//p[@itemprop = "description"]')))
                                offer_description_sections = driver.find_elements_by_xpath('//div[@class = "left overflow-hidden agent-content"]//div[@class = "sect "]//div[@class = "sect "]|//p[@itemprop = "description"]')
                                for piece in offer_description_sections:
                                    total_string += piece.text + "\n"
                            except TimeoutException:
                                driver.refresh()
                            except Exception as e:
                                print(str(e)+offer_URL)
                                pass
                            if any(x in total_string.lower() for x in ignore_string):
                                break_marker = True
                                print("FOUND NON-STANDARD CONSTRUCTION OR LEASEHOLD: " + offer_URL)
                            for group in separate_file_string:
                                for element in group:
                                    if element in total_string.lower():
                                        output_file_picker = separate_file_string.index(group)
                            no_of_beds = WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//h1[@class = "fs-22"]'))).text.split()[0]
                            house_asking_price = int("".join(driver.find_element_by_xpath('//p[@id = "propertyHeaderPrice"][@class = "property-header-price "]//strong').text.split("£")[1].split(",")))
                            #####Click on the broadband search to get the postcode
                            WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.XPATH, '//span[@class = "check-broadband-speed"]'))).click()
                            WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.XPATH, '//a[@class = "see-all-offers"]'))).click()
                        except ElementClickInterceptedException:
                            driver.refresh()
                        except IndexError:
                            #####If when the house price string is split an error is caused by expectancy nonconformace it will do the below check and then move to the next offer
                            if driver.find_element_by_xpath('//p[@id = "propertyHeaderPrice"][@class = "property-header-price "]//strong').text == "POA":
                                break_marker = True
                                break
                            else:
                                driver.refresh()
                        except TimeoutException:
                            driver.refresh()
                        else:
                            break
                    if break_marker == True:
                        try:
                            driver.switch_to.window(driver.window_handles[1])
                            driver.close()
                            driver.switch_to.window(driver.window_handles[0])
                        except IndexError:
                            pass
                        break
                    while True:
                        try:
                            #####Switches to the tab opened by the broadband search goign to Compare the market
                            driver.switch_to.window(driver.window_handles[1])
                        except IndexError:
                            WebDriverWait(driver, 2).until(EC.element_to_be_clickable((By.XPATH, '//a[@class = "see-all-offers"]'))).click()
                        else:
                            break
                    infinite_loop_police = 0
                    while True:
                        try:
                            post_code = WebDriverWait(driver, 5).until(EC.visibility_of_element_located((By.XPATH,'//span[@class = "current-provider-filter__text"][@id = "current-provider-filter__location-text"]'))).text
                        except ElementClickInterceptedException:
                            driver.refresh()
                        except TimeoutException:
                            if infinite_loop_police < 6:
                                driver.refresh()
                                infinite_loop_police += 1
                            else:
                                print("Could not load compare the market for offer: " + offer_URL)
                                driver.close()
                                driver.switch_to.window(driver.window_handles[0])
                                break_marker = True
                                break
                        else:
                            break

                    if break_marker == True:
                        break

                    #####Closes the broadband search and moved back to the main tab
                    driver.close()
                    driver.switch_to.window(driver.window_handles[0])
                    #####Prepares to search for rents starting with 1/4 mile radius then 1/2
                    radius = "0.25"
                    driver.get('https://www.rightmove.co.uk/property-to-rent/search.html?searchLocation=' + post_code.split()[0] + '+' + post_code.split()[1] + '&locationIdentifier=&useLocationIdentifier=false&rent=To+rent')
                    while True:
                        try:
                            #####Configures the search with the number of beds of our sale offer and other standard parameters
                            WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH,'//select[@id = "radius"][@name = "radius"]//option[@value = "'+radius+'"]'))).click()
                            driver.find_element_by_xpath('//select[@id = "minBedrooms"][@name = "minBedrooms"]//option[@value = "' + no_of_beds + '"]').click()
                            driver.find_element_by_xpath('//select[@id = "maxBedrooms"][@name = "maxBedrooms"]//option[@value = "' + no_of_beds + '"]').click()
                            driver.find_element_by_xpath('//span[@class = "tickbox--indicator"]').click()
                            driver.find_element_by_xpath('//select[@id = "displayPropertyType"][@name = "displayPropertyType"]//option[@value = "houses"]').click()
                            driver.find_element_by_xpath('//button[@id = "submit"][@class = "button touchsearch-button touchsearch-primarybutton"]').click()
                        except TimeoutException:
                            driver.refresh()
                        else:
                            break
                    driver.get(driver.current_url + "&dontShow=retirement")
                    #####Initializes values necessary for ROI calculations below
                    total_rent = 0
                    average_rent = 0
                    offer_percentage_ROI = 0
                    while True:
                        try:
                            #####Grabs a list of all the rent offers and loops through all of them to add them upp and get the average
                            WebDriverWait(driver, 2).until(EC.visibility_of_element_located((By.XPATH, '//span[@class = "propertyCard-priceValue"]')))
                            rent_search_URL = str(driver.current_url)
                            rents_list = driver.find_elements_by_xpath('//span[@class = "propertyCard-priceValue"]')
                            for rent_offer in rents_list:
                                new_rent_found = rent_offer.text.split()[0][1:]
                                if "," in new_rent_found:
                                    new_rent_found = new_rent_found.replace(",", "")
                                total_rent = total_rent + int(new_rent_found)
                            average_rent = total_rent / len(rents_list)
                        except TimeoutException:
                            try:
                                if driver.find_element_by_xpath('//p[@class = "enhancedZeroResults-title"]').text:
                                    if radius == "0.25":
                                        radius = "0.5"
                                        driver.find_element_by_xpath('//select[@class = "select"][@name = "radius"]//option[@value = "' + radius + '"]').click()
                                    else:
                                        break
                            except NoSuchElementException:
                                driver.refresh()
                                print("Illogical index caught")
                        except IndexError:
                            driver.refresh()
                        else:
                            break
                    #####Calculates the ROI
                    if not BRRR and average_rent != 0:
                        offer_yearly_profit = average_rent * 12 * (50/52) * 0.8 - house_asking_price * 0.75 * 0.03
                        offer_percentage_ROI = int((offer_yearly_profit / (house_asking_price *0.3)) * 100)
                    elif BRRR:
                        print("Still to be coded for BRRR")
                    #####Checks ROI is within what is desired
                    if offer_percentage_ROI != 0 and offer_percentage_ROI >= target_ROI:
                        if output_file_picker == "":
                            if os.path.isfile(Path(work_area_path, "findings.docx")):
                                output_file = docx.Document(Path(work_area_path, "findings.docx"))
                            else:
                                output_file = docx.Document()
                            if trailing_list[0] == 0:
                                if type(sales_ratio) == float and type(rent_ratio) == float:
                                    output_file.add_paragraph(area_of_search + " - area code " + code + " - sales supply/demand is " + str(round(sales_ratio, 2)) + " - rent supply/demand is " + str(round(rent_ratio, 2)))
                                elif type(sales_ratio) == str:
                                    output_file.add_paragraph(area_of_search + " - area code " + code + " - " + str(sales_ratio) + " sales" + " - rent supply/demand is " + str(round(rent_ratio, 2)))
                                elif type(rent_ratio) == str:
                                    output_file.add_paragraph(area_of_search + " - area code " + code + " - sales supply/demand is " + str(round(sales_ratio, 2)) + " - " + str(rent_ratio) + "rent")
                                paragraph = output_file.add_paragraph(str(offer_percentage_ROI) + "%, £" + str(house_asking_price) + ", £" + str(round(int(average_rent), 0)) + ", " + radius + " miles, ")
                                add_hyperlink(paragraph, "Offer URL", offer_URL)
                                paragraph.add_run(", ")
                                add_hyperlink(paragraph, "Rent search URL", rent_search_URL)
                                paragraph.add_run().add_break()
                                trailing_list[0] = 1
                            else:
                                all_doc_paragraphs = output_file.paragraphs
                                paragraph = all_doc_paragraphs[-1]
                                paragraph.add_run(str(offer_percentage_ROI) + "%, £" + str(house_asking_price) + ", £" + str(round(int(average_rent),0)) + ", " + radius + " miles, ")
                                add_hyperlink(paragraph, "Offer URL", offer_URL)
                                paragraph.add_run(", ")
                                add_hyperlink(paragraph, "Rent search URL", rent_search_URL)
                                paragraph.add_run().add_break()
                            output_file.save(Path(work_area_path,"findings.docx"))
                        else:
                            if os.path.isfile(Path(work_area_path, separate_file_string[output_file_picker][0] + ".docx")):
                                separate_output_file = docx.Document(Path(work_area_path, separate_file_string[output_file_picker][0] + ".docx"))
                            else:
                                separate_output_file = docx.Document()
                            if trailing_list[output_file_picker+1] == 0 :
                                if type(sales_ratio) == float and type(rent_ratio) == float:
                                    separate_output_file.add_paragraph(area_of_search + " - area code " + code + " - sales supply/demand is " + str(round(sales_ratio, 2)) + " - rent supply/demand is " + str(round(rent_ratio, 2)))
                                elif type(sales_ratio) == str:
                                    separate_output_file.add_paragraph(area_of_search + " - area code " + code + " - " + str(sales_ratio) + " sales" + " - rent supply/demand is " + str(round(rent_ratio, 2)))
                                elif type(rent_ratio) == str:
                                    separate_output_file.add_paragraph(area_of_search + " - area code " + code + " - sales supply/demand is " + str(round(sales_ratio, 2)) + " - " + str(rent_ratio) + "rent")
                                separate_paragraph = separate_output_file.add_paragraph(str(offer_percentage_ROI) + "%, £" + str(house_asking_price) + ", £" + str(round(int(average_rent), 0)) + ", " + radius + " miles, ")
                                add_hyperlink(separate_paragraph, "Offer URL", offer_URL)
                                separate_paragraph.add_run(", ")
                                add_hyperlink(separate_paragraph, "Rent search URL", rent_search_URL)
                                separate_paragraph.add_run().add_break()
                                trailing_list[output_file_picker+1] = 1
                            else:
                                all_doc_paragraphs = separate_output_file.paragraphs
                                separate_paragraph = all_doc_paragraphs[-1]
                                separate_paragraph.add_run(str(offer_percentage_ROI) + "%, £" + str(house_asking_price) + ", £" + str(round(int(average_rent), 0)) + ", " + radius + " miles, ")
                                add_hyperlink(separate_paragraph, "Offer URL", offer_URL)
                                separate_paragraph.add_run(", ")
                                add_hyperlink(separate_paragraph, "Rent search URL", rent_search_URL)
                                separate_paragraph.add_run().add_break()
                            separate_output_file.save(Path(work_area_path, separate_file_string[output_file_picker][0] + ".docx"))
                    # driver.execute_script("window.open('https://www.google.com', 'new window')")
                except WebDriverException:
                    driver = webdriver.Chrome("C:/Users/Cristi/AppData/Local/Programs/Python/CALC Investing/ChromeDriver/chromedriver.exe",desired_capabilities=Page_load_strategy)
                    driver.maximize_window()
                    print("Unexpected page crash")
                else:
                    break
            lock.acquire()
            if current_offer_in_page_increment < number_of_offers_for_page:
                current_offer_in_page_increment += 1
            lock.release()
        lock.acquire()
        current_page_number += 1
        lock.release()

stop_marker = True
thread.join()

end_time = datetime.now()
script_exec_time = end_time - begin_time
output_file = docx.Document(Path(work_area_path, "findings.docx"))
output_file.add_paragraph("Run time: " + str(script_exec_time) + "; " + str(offer_counter) + " properties were checked")
output_file.save(Path(work_area_path, "findings.docx"))